import streamlit as st
import requests
from pytube import YouTube
import tempfile
import shutil
import os
from fpdf import FPDF
from docx import Document
import fitz  # PyMuPDF
from groq import Groq

# Function to get a temporary directory
def get_temp_dir():
    return tempfile.mkdtemp()


# Function to download YouTube audio
def download_youtube_audio(youtube_url, temp_dir):
    try:
        yt = YouTube(youtube_url)
        stream = yt.streams.filter(only_audio=True).first()
        if not stream:
            raise ValueError("No audio stream found in the YouTube video.")
        temp_file_path = os.path.join(temp_dir, "audio.mp4")
        stream.download(output_path=temp_dir, filename="audio.mp4")
        return temp_file_path
    except Exception as e:
        st.error(f"Error downloading YouTube video: {e}")
        return None


# Function to clean up temporary files
def clean_up_temp_files(temp_dir):
    shutil.rmtree(temp_dir)


# Function to render custom styles for the app
def render_custom_styles():
    st.markdown(
        """
        <style>
        .generated-content {
            height: 600px !important; /* Increase height to 600 pixels */
            width: 100%; /* Set width to 100% */
            overflow-y: auto; /* Enable vertical scroll if content exceeds height */
            border: 1px solid #ccc; /* Optional: Add border for visibility */
            padding: 10px; /* Optional: Add padding for better readability */
            margin-bottom: 20px; /* Optional: Add margin for spacing between content */
        }
        </style>
        """,
        unsafe_allow_html=True
    )


# Function to transcribe audio using Deepgram API
def transcribe_audio_deepgram(audio_path, api_key):
    try:
        with open(audio_path, "rb") as audio_file:
            response = requests.post("https://api.deepgram.com/v1/listen", headers={"Authorization": f"Token {api_key}", "Content-Type": "audio/wav"}, data=audio_file)
            response.raise_for_status()
            transcription = response.json().get('results', {}).get('channels', [])[0].get('alternatives', [])[0].get('transcript')
            return transcription
    except Exception as e:
        st.error(f"Failed to transcribe audio: {e}")
        return ""


# Function to handle transcription from audio file or YouTube URL
def transcribe_audio_and_get_transcription(audio_file, youtube_url, deepgram_api_key):
    temp_dir = get_temp_dir()

    if audio_file:
        audio_path = os.path.join(temp_dir, f"temp_audio.{audio_file.name.split('.')[-1]}")
        with open(audio_path, "wb") as f:
            f.write(audio_file.getbuffer())
        transcription_output = transcribe_audio_deepgram(audio_path, deepgram_api_key)

    elif youtube_url:
        audio_path = download_youtube_audio(youtube_url, temp_dir)
        if audio_path:
            transcription_output = transcribe_audio_deepgram(audio_path, deepgram_api_key)

    clean_up_temp_files(temp_dir)
    return transcription_output


# Function to extract text from PDF files
def extract_text_from_pdf(file):
    text = ""
    with fitz.open(stream=file.read(), filetype="pdf") as doc:
        for page in doc:
            text += page.get_text()
    return text


# Function to extract text from DOCX files
def extract_text_from_docx(file):
    doc = Document(file)
    return "\n".join([para.text for para in doc.paragraphs])


# Function to generate lecture notes using Groq API
def generate_notes(transcription, api_key, lesson_plan_text=None):
    client = Groq(api_key=api_key)
    prompt = f"""Create detailed lecture notes summarizing the key concepts discussed in the provided transcription. Highlight important topics and keep the notes concise and organized.
    After the notes, create a structured document with the discussed topics and associated timestamps from the original transcription. Ensure the timestamps accurately reflect the timing of each topic's discussion:\n\n{transcription}"""
    if lesson_plan_text:
        prompt += f"\n\nPlease ensure the notes align with the following lesson plan:\n\n{lesson_plan_text}"
    response = client.chat.completions.create(
        model="mixtral-8x7b-32768",
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content.strip()


# Function to generate quiz using Groq API
def generate_quiz(transcription, api_key, num_questions):
    client = Groq(api_key=api_key)
    prompt = f"Generate {num_questions} multiple-choice questions with answers given at the end for all the questions, keep a balance of easy, moderate and difficult questions from the following transcription:\n\n{transcription}"
    response = client.chat.completions.create(
        model="mixtral-8x7b-32768",
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content.strip()


# Function to create a PDF document
def create_pdf(text, file_name):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    text = text.encode('latin-1', 'replace').decode('latin-1')
    pdf.multi_cell(0, 10, text)
    pdf.output(file_name)


# Function to create a Word document
def create_word_doc(text, file_name):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(file_name)


# Function to render the homepage
def render_homepage():
    st.markdown("<h1 style='font-family:Georgia; font-size:36px;'>AI-Powered Teaching Assistant</h1>", unsafe_allow_html=True)
    st.write("<p style='font-family:Arial; font-size:18px;'>Welcome to the AI-powered teaching assistant. This tool helps you generate lecture notes and quizzes from audio recordings and YouTube videos.</p>", unsafe_allow_html=True)

    st.header("API Keys")
    deepgram_api_key = st.text_input("Enter your Deepgram API key:")
    groq_api_key = st.text_input("Enter your Groq API key:")

    if st.button("Save API Keys"):
        st.session_state["deepgram_api_key"] = deepgram_api_key
        st.session_state["groq_api_key"] = groq_api_key
        st.success("API keys saved for this session.")


# Function to render inputs for audio file or YouTube URL
def render_inputs(key):
    st.markdown("<h2 style='font-family:Georgia; color:#2C3E50; font-size:24px;'>Upload Lecture Recording or Enter YouTube Video URL</h2>", unsafe_allow_html=True)
    st.markdown("<p style='font-family:Arial; font-size:16px; color:#34495E;'>Upload your lecture audio file:</p>", unsafe_allow_html=True)
    audio_file = st.file_uploader("", type=["mp3", "wav", "m4a"], key=f"{key}_audio")
    st.markdown("<p style='font-family:Arial; font-size:16px; color:#34495E;'>Paste YouTube video URL here:</p>", unsafe_allow_html=True)
    youtube_url = st.text_input("", key="notes_youtube_url")
    return audio_file, youtube_url


# Function to process audio file or YouTube URL and generate notes or quiz
def process_audio(audio_file, youtube_url, generate_func, num_questions=None, lesson_plan_text=None):
    temp_dir = get_temp_dir()

    if audio_file:
        st.success("Processing the audio file...")
        audio_path = os.path.join(temp_dir, f"temp_audio.{audio_file.name.split('.')[-1]}")
        with open(audio_path, "wb") as f:
            f.write(audio_file.getbuffer())
        transcription_output = transcribe_audio_deepgram(audio_path, st.session_state["deepgram_api_key"])

    elif youtube_url:
        st.success("Processing the YouTube video...")
        audio_path = download_youtube_audio(youtube_url, temp_dir)
        if audio_path:
            transcription_output = transcribe_audio_deepgram(audio_path, st.session_state["deepgram_api_key"])

    if transcription_output:
        if generate_func.__name__ == "generate_notes":
            output = generate_func(transcription_output, st.session_state["groq_api_key"], lesson_plan_text)
        else:
            output = generate_func(transcription_output, st.session_state["groq_api_key"], num_questions)
        st.markdown(f'<div class="generated-content">{output}</div>', unsafe_allow_html=True)
        render_download_options(output, generate_func.__name__)

    clean_up_temp_files(temp_dir)


# Function to render download options for generated content
def render_download_options(output, output_type):
    #st.download_button(f"Download {output_type} as TXT", output, file_name=f"{output_type}.txt")
    pdf_file = f"{output_type}.pdf"
    create_pdf(output, pdf_file)
    with open(pdf_file, "rb") as f:
        st.download_button(f"Download {output_type} as PDF", f, file_name=pdf_file)
    word_file = f"{output_type}.docx"
    create_word_doc(output, word_file)
    with open(word_file, "rb") as f:
        st.download_button(f"Download {output_type} as Word Document", f, file_name=word_file)


# Function to render the notes and quiz generation page
def render_notes_and_quiz_page():
    st.title("Notes and Quiz Generation")

    col1, col2 = st.columns(2)
    with col1:
        st.markdown("<p style='font-family:Arial; font-size:16px;'>Upload your lecture audio file:</p>", unsafe_allow_html=True)
        audio_file = st.file_uploader("", type=["mp3", "wav", "m4a"], key=f"notes_quiz_audio")
    with col2:
        st.markdown("<p style='font-family:Arial; font-size:16px;'>Paste YouTube video URL here:</p>", unsafe_allow_html=True)
        youtube_url = st.text_input("", key="notes_quiz_youtube_url")

    st.markdown("<h3 style='font-family:Georgia; font-size:20px;'>Upload Lesson Plan (Optional)</h3>", unsafe_allow_html=True)
    lesson_plan_file = st.file_uploader("Upload your lesson plan document", type=["pdf", "docx", "txt"], key="notes_quiz_lesson_plan")

    lesson_plan_text = None
    if lesson_plan_file:
        if lesson_plan_file.name.endswith('.pdf'):
            lesson_plan_text = extract_text_from_pdf(lesson_plan_file)
        elif lesson_plan_file.name.endswith('.docx'):
            lesson_plan_text = extract_text_from_docx(lesson_plan_file)
        else:
            st.error("Unsupported file type. Please upload a PDF or DOCX file.")

    num_questions = st.number_input("Number of quiz questions", min_value=5, max_value=20, value=10, step=1, key="notes_quiz_num_questions")

    col1, col2, col3 = st.columns(3)
    with col1:
        if st.button("Generate Notes"):
            # Clear previous downloads
            clear_session_state_downloads()
            process_audio(audio_file, youtube_url, generate_notes, lesson_plan_text=lesson_plan_text)

    with col2:
        if st.button("Generate Quiz"):
            # Clear previous downloads
            clear_session_state_downloads()
            process_audio(audio_file, youtube_url, generate_quiz, num_questions=int(num_questions))

    with col3:
        if st.button("Generate Notes and Quiz"):
            # Clear previous downloads
            clear_session_state_downloads()
            if audio_file or youtube_url:
                combined_transcription = transcribe_audio_and_get_transcription(audio_file, youtube_url, st.session_state["deepgram_api_key"])
                if combined_transcription:
                    notes_output = generate_notes(combined_transcription, st.session_state["groq_api_key"], lesson_plan_text)
                    st.session_state["notes_output"] = notes_output  # Store notes output in session state
                    st.markdown(f'<div class="generated-content">{notes_output}</div>', unsafe_allow_html=True)

                    quiz_output = generate_quiz(combined_transcription, st.session_state["groq_api_key"], int(num_questions))
                    st.session_state["quiz_output"] = quiz_output  # Store quiz output in session state
                    st.markdown(f'<div class="generated-content">{quiz_output}</div>', unsafe_allow_html=True)
            else:
                st.error("Please upload an audio file or enter a YouTube URL.")


    # Render download buttons based on session state
    if "notes_output" in st.session_state:
        render_download_options(st.session_state["notes_output"], "Notes")

    if "quiz_output" in st.session_state:
        render_download_options(st.session_state["quiz_output"], "Quiz")


    # Apply styles to the buttons
    st.markdown("""
        <style>
        div.stButton > button {
            width: 100%;
            font-size: 20px;
            padding: 10px;
            border-radius: 5px;
            background-color: #2C3E50;
            color: white;
            font-weight: bold;
            margin-top: 20px;
        }
        div.stButton > button:hover {
            background-color: #34495E;
        }
        </style>
    """, unsafe_allow_html=True)

def clear_session_state_downloads():
    # Clear previously stored outputs in session state
    if "notes_output" in st.session_state:
        del st.session_state["notes_output"]

    if "quiz_output" in st.session_state:
        del st.session_state["quiz_output"]

def render_footer():
    st.markdown("---")
    st.write("<p style='font-family:Arial; font-size:14px; color:#7F8C8D;'>© 2024 AI-Powered Teaching Assistant</p>", unsafe_allow_html=True)



def main():
    st.set_page_config(page_title="AI-Powered Teaching Assistant", page_icon=":books:", layout="wide")
    st.markdown("<style>body { background-color: #FFFFFF; }</style>", unsafe_allow_html=True)
    render_custom_styles()
    st.sidebar.title("Navigation")
    page = st.sidebar.radio("Go to", ["Home", "Notes and Quiz Generation"])

    if page == "Home":
        render_homepage()
    elif page == "Notes and Quiz Generation":
        render_notes_and_quiz_page()

    render_footer()

if __name__ == "__main__":
    main()