import streamlit as st
import requests
from pytube import YouTube
import tempfile
import shutil
import os
from fpdf import FPDF
from docx import Document
import fitz
from groq import Groq

DEEPGRAM_API_KEY = "3963c922b8b16ac85c45a673ab42d6e842151e9a"
GROQ_API_KEY = "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"

def get_temp_dir():
    return tempfile.mkdtemp()

def reset_inputs():
    st.session_state.clear()

def download_youtube_audio(youtube_url, temp_dir):
    try:
        yt = YouTube(youtube_url)
        stream = yt.streams.filter(only_audio=True).first()
        if not stream:
            raise ValueError("No audio stream found in the YouTube video.")
        temp_file_path = os.path.join(temp_dir, "audio.mp4")
        stream.download(output_path=temp_dir, filename="audio.mp4")
        return temp_file_path
    except Exception as e:
        st.error(f"Error downloading YouTube video: {e}")
        return None

def clean_up_temp_files(temp_dir):
    shutil.rmtree(temp_dir)

def transcribe_audio_deepgram(audio_path, api_key):
    api_url = "https://api.deepgram.com/v1/listen"
    headers = {
        "Authorization": f"Token {api_key}",
        "Content-Type": "audio/wav"
    }

    try:
        with open(audio_path, "rb") as audio_file:
            response = requests.post(api_url, headers=headers, data=audio_file)
        response.raise_for_status()
        transcription = response.json().get('results', {}).get('channels', [])[0].get('alternatives', [])[0].get('transcript')
        return transcription
    except Exception as e:
        st.error(f"Failed to transcribe audio: {e}")
        return ""

def transcribe_audio_and_get_transcription(audio_file, youtube_url, deepgram_api_key):
    temp_dir = get_temp_dir()
    transcription_output = None

    if audio_file:
        audio_path = os.path.join(temp_dir, f"temp_audio.{audio_file.name.split('.')[-1]}")
        with open(audio_path, "wb") as f:
            f.write(audio_file.getbuffer())
        transcription_output = transcribe_audio_deepgram(audio_path, deepgram_api_key)

    elif youtube_url:
        audio_path = download_youtube_audio(youtube_url, temp_dir)
        if audio_path:
            transcription_output = transcribe_audio_deepgram(audio_path, deepgram_api_key)

    clean_up_temp_files(temp_dir)
    return transcription_output

def extract_text_from_pdf(file):
    text = ""
    with fitz.open(stream=file.read(), filetype="pdf") as doc:
        for page in doc:
            text += page.get_text()
    return text

def extract_text_from_docx(file):
    doc = Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def generate_notes(transcription, api_key, lesson_plan_text=None):
    client = Groq(api_key=api_key)
    prompt = f"Generate detailed lecture notes and after the notes at the end of the notes, generate a structured document containing discussed topics with associated timestamps from the following transcription:\n\n{transcription}"
    if lesson_plan_text:
        prompt += f"\n\nPlease ensure the notes align with the following lesson plan:\n\n{lesson_plan_text}"
    response = client.chat.completions.create(
        model="mixtral-8x7b-32768",
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content.strip()

def generate_quiz(transcription, api_key, num_questions):
    client = Groq(api_key=api_key)
    prompt = f"Generate {num_questions} multiple-choice questions with answers given at the end for all the questions from the following transcription:\n\n{transcription}"
    response = client.chat.completions.create(
        model="mixtral-8x7b-32768",
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content.strip()

def create_pdf(text, file_name):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    text = text.encode('latin-1', 'replace').decode('latin-1')
    pdf.multi_cell(0, 10, text)
    pdf.output(file_name)

def create_word_doc(text, file_name):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(file_name)

def render_homepage():
    st.markdown("<h1 style='font-family:Georgia; color:#2C3E50; font-size:36px;'>AI-Powered Teaching Assistant</h1>", unsafe_allow_html=True)
    st.write("<p style='font-family:Arial; font-size:18px; color:#34495E;'>Welcome to the AI-powered teaching assistant. This tool helps you generate lecture notes and quizzes from audio recordings and YouTube videos.</p>", unsafe_allow_html=True)

def render_inputs(key):
    st.markdown("<h2 style='font-family:Georgia; color:#2C3E50; font-size:24px;'>Upload Lecture Recording or Enter YouTube Video URL</h2>", unsafe_allow_html=True)
    st.markdown("<p style='font-family:Arial; font-size:16px; color:#34495E;'>Upload your lecture audio file:</p>", unsafe_allow_html=True)
    audio_file = st.file_uploader("", type=["mp3", "wav", "m4a"], key=f"{key}_audio")
    st.markdown("<p style='font-family:Arial; font-size:16px; color:#34495E;'>Paste YouTube video URL here:</p>", unsafe_allow_html=True)
    youtube_url = st.text_input("", key="notes_youtube_url")
    return audio_file, youtube_url

def process_audio(audio_file, youtube_url, generate_func, num_questions=None, lesson_plan_text=None):
    temp_dir = get_temp_dir()

    if audio_file:
        st.success("Processing the audio file...")
        audio_path = os.path.join(temp_dir, f"temp_audio.{audio_file.name.split('.')[-1]}")
        with open(audio_path, "wb") as f:
            f.write(audio_file.getbuffer())
        transcription_output = transcribe_audio_deepgram(audio_path, DEEPGRAM_API_KEY)
        if transcription_output:
            if generate_func.__name__ == "generate_notes":
                output = generate_func(transcription_output, GROQ_API_KEY, lesson_plan_text)
            else:
                output = generate_func(transcription_output, GROQ_API_KEY, num_questions)
            st.text_area("Generated Output", output, height=300)
            render_download_options(output, generate_func.__name__)

    elif youtube_url:
        st.success("Processing the YouTube video...")
        audio_path = download_youtube_audio(youtube_url, temp_dir)
        if audio_path:
            transcription_output = transcribe_audio_deepgram(audio_path, DEEPGRAM_API_KEY)
            if transcription_output:
                if generate_func.__name__ == "generate_notes":
                    output = generate_func(transcription_output, GROQ_API_KEY, lesson_plan_text)
                else:
                    output = generate_func(transcription_output, GROQ_API_KEY, num_questions)
                st.text_area("Generated Output", output, height=300)
                render_download_options(output, generate_func.__name__)
        else:
            st.error("Failed to download YouTube video.")
    else:
        st.error("Please upload an audio file or enter a YouTube URL.")

    clean_up_temp_files(temp_dir)

def render_download_options(output, output_type):
    st.download_button(f"Download {output_type} as TXT", output, file_name=f"{output_type}.txt")
    pdf_file = f"{output_type}.pdf"
    create_pdf(output, pdf_file)
    with open(pdf_file, "rb") as f:
        st.download_button(f"Download {output_type} as PDF", f, file_name=pdf_file)
    word_file = f"{output_type}.docx"
    create_word_doc(output, word_file)
    with open(word_file, "rb") as f:
        st.download_button(f"Download {output_type} as Word Document", f, file_name=word_file)

def render_lecture_notes_page():
    st.title("Lecture Notes Generation")

    # Input section
    #Sst.markdown("<h3 style='font-family:Georgia; font-size:20px;'>Upload Lecture Recording or Enter YouTube Video URL</h3>", unsafe_allow_html=True)
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("<p style='font-family:Arial; font-size:16px;'>Upload your lecture audio file:</p>", unsafe_allow_html=True)
        audio_file = st.file_uploader("", type=["mp3", "wav", "m4a"], key=f"notes_audio")
    with col2:
        st.markdown("<p style='font-family:Arial; font-size:16px;'>Paste YouTube video URL here:</p>", unsafe_allow_html=True)
        youtube_url = st.text_input("", key="notes_youtube_url")

    # Settings section
    st.markdown("<h3 style='font-family:Georgia; font-size:20px;'>Upload Lesson Plan (Optional)</h3>", unsafe_allow_html=True)
    lesson_plan_file = st.file_uploader("Upload your lesson plan document", type=["pdf", "docx", "txt"], key="notes_lesson_plan")

    # Transcription and Notes generation section
    lesson_plan_text = None
    if lesson_plan_file:
        if lesson_plan_file.name.endswith('.pdf'):
            lesson_plan_text = extract_text_from_pdf(lesson_plan_file)
        elif lesson_plan_file.name.endswith('.docx'):
            lesson_plan_text = extract_text_from_docx(lesson_plan_file)
        else:
            st.error("Unsupported file type. Please upload a PDF or DOCX file.")

    if st.button("Generate Notes"):
        process_audio(audio_file, youtube_url, generate_notes, lesson_plan_text=lesson_plan_text)

    st.markdown("<h3 style='font-family:Georgia; font-size:20px;'>Generate Notes and Quiz</h3>", unsafe_allow_html=True)
    num_questions = st.number_input("Number of quiz questions", min_value=5, max_value=20, value=10, step=1, key="combined_num_questions")

    if st.button("Generate"):
        if audio_file or youtube_url:
            combined_transcription = transcribe_audio_and_get_transcription(audio_file, youtube_url, DEEPGRAM_API_KEY)
            if combined_transcription:
                notes_output = generate_notes(combined_transcription, GROQ_API_KEY, lesson_plan_text)
                st.text_area("Generated Notes", notes_output, height=300)
                render_download_options(notes_output, "notes")

                quiz_output = generate_quiz(combined_transcription, GROQ_API_KEY, int(num_questions))
                st.text_area("Generated Quiz", quiz_output, height=300)
                render_download_options(quiz_output, "quiz")
        else:
            st.error("Please upload an audio file or enter a YouTube URL.")

def render_quiz_generation_page():
    st.title("Quiz Generation")
    st.markdown("<h3 style='font-family:Georgia;  font-size:20px;'>Upload Lecture Recording or Enter YouTube Video URL</h3>", unsafe_allow_html=True)
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("<p style='font-family:Arial; font-size:16px;'>Upload your lecture audio file:</p>", unsafe_allow_html=True)
        audio_file = st.file_uploader("", type=["mp3", "wav", "m4a"], key=f"notes_audio")
    with col2:
        st.markdown("<p style='font-family:Arial; font-size:16px;'>Paste YouTube video URL here:</p>", unsafe_allow_html=True)
        youtube_url = st.text_input("", key="notes_youtube_url")



    st.header("Settings")
    num_questions = st.slider("Number of questions to generate", min_value=5, max_value=20, value=10, key="quiz_num_questions")
    if st.button("Generate Quiz"):
        process_audio(audio_file, youtube_url, generate_quiz, num_questions)
    
def render_footer():
    st.markdown("---")
    st.write("<p style='font-family:Arial; font-size:14px; color:#7F8C8D;'>© 2024 AI-Powered Teaching Assistant</p>", unsafe_allow_html=True)

def main():
    st.set_page_config(page_title="AI-Powered Teaching Assistant", page_icon=":books:", layout="wide")
    st.markdown("<style>body { background-color: #FFFFFF; }</style>", unsafe_allow_html=True)
    st.sidebar.title("Navigation")
    page = st.sidebar.radio("Go to", ["Home", "Lecture Notes", "Quiz Generation"])

    if page == "Home":
        render_homepage()
    elif page == "Lecture Notes":
        render_lecture_notes_page()
    elif page == "Quiz Generation":
        render_quiz_generation_page()

    #render_footer()

if __name__ == "__main__":
    main()